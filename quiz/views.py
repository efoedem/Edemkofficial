import os
import csv
import io
import json
import random
from django.contrib.messages import get_messages
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import user_passes_test
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt  # 🛡️ Core bypass helper
from .models import Lecturer, Course, Question, StudentSubmission, AllowedStudent




try:
    import openpyxl
except ImportError:
    openpyxl = None

from docx import Document
from pypdf import PdfReader


# ==========================================================
#             STUDENT VERIFICATION & PORTAL VIEWS
# ==========================================================


from math import radians, cos, sin, asin, sqrt


def haversine(lon1, lat1, course):
    # CORRECTED: Using the field names defined in your models.py
    # course.latitude and course.longitude
    exam_lat = float(course.latitude)
    exam_lng = float(course.longitude)

    # Convert to radians
    lon1, lat1, exam_lng, exam_lat = map(radians, [lon1, lat1, exam_lng, exam_lat])

    dlon = exam_lng - lon1
    dlat = exam_lat - lat1
    a = sin(dlat / 2) ** 2 + cos(lat1) * cos(exam_lat) * sin(dlon / 2) ** 2
    c = 2 * asin(sqrt(a))
    r = 6371000  # Earth radius in meters
    return c * r


@csrf_exempt
def login_portal(request):
    if request.method == "POST":
        course_id = request.POST.get("course_id")
        index_number = request.POST.get("index_number", "").strip().upper()

        user_lat = float(request.POST.get("lat", 0))
        user_lng = float(request.POST.get("lng", 0))

        try:
            course = get_object_or_404(Course, id=course_id)

            # 1. Location Check
            distance = haversine(user_lng, user_lat, course)
            if distance > course.radius_meters:
                messages.error(request, f"Security Lockdown: Outside exam zone (Detected: {int(distance)}m).")
                return redirect("login_portal")

            # 2. Time Check
            now = timezone.now()
            if now < course.start_time:
                messages.error(request, f"Examination begins at {course.start_time.strftime('%I:%M %p')}.")
                return redirect("login_portal")
            if now > course.end_time:
                messages.error(request, "Examination window has closed.")
                return redirect("login_portal")

            full_name = request.POST.get("full_name", "Anonymous Student").strip()

            request.session['is_authenticated'] = True
            request.session['student_name'] = full_name
            request.session['index_number'] = index_number
            request.session['course_id'] = course.id

            return redirect('start_quiz')

        except Exception as e:
            messages.error(request, f"Authentication error: {str(e)}")
            return redirect("login_portal")

    # GET Request Logic
    # ... (Keep your existing GET logic)

    # GET Request Logic
    lecturers = Lecturer.objects.all()
    context = {
        'lecturers': lecturers,
        'is_authenticated': request.session.get('is_authenticated', False),
        'student_name': request.session.get('student_name', ''),
        'index_number': request.session.get('index_number', ''),
    }

    if request.session.get('course_id'):
        course = Course.objects.filter(id=request.session.get('course_id')).first()
        context['course'] = course
        now = timezone.now()
        # Pass status to help display buttons/messages in login.html
        context['status'] = 'WAITING' if timezone.now() < course.start_time else 'ACTIVE'

    return render(request, 'quiz/login.html', context)


def get_courses(request):
    """
    Dynamically returns JSON data filtered by lecturer selection
    to populate the frontend course selection field.
    """
    lecturer_id = request.GET.get('lecturer_id')
    if lecturer_id:
        courses = Course.objects.filter(lecturer_id=lecturer_id).values('id', 'code', 'title')
        return JsonResponse(list(courses), safe=False)
    return JsonResponse([], safe=False)


def logout_portal(request):
    request.session.flush() # This clears everything, allowing them to login again
    return redirect('login_portal')



def submit_assignment(request):
    if request.method == 'POST':
        # Get the uploaded file
        file = request.FILES.get('file')

        # Create the submission instance
        submission = StudentSubmission(
            student_name=request.POST.get('student_name'),
            index_number=request.POST.get('index_number'),
            course_id=request.POST.get('course_id'),
            submission_type='ASSIGNMENT',
            # Read file into binary data and store name
            file_data=file.read() if file else None,
            file_name=file.name if file else None
        )
        submission.save()

        messages.success(request, "Assignment submitted successfully!")
        return render(request, 'quiz/submitted.html')  # Ensure you have this template

    # Only show courses that are 'ASSIGNMENT' type in the dropdown
    courses = Course.objects.filter(assessment_type='ASSIGNMENT')
    return render(request, 'quiz/submit.html', {'courses': courses})


@csrf_exempt
def start_quiz(request):
    # 1. Validate Session Existence
    print(f"Session Data: {request.session.items()}")
    student_name = request.session.get('student_name')
    index_number = request.session.get('index_number')
    course_id = request.session.get('course_id')

    if not student_name or not index_number or not course_id:
        messages.error(request, "Session expired. Please login again.")
        return redirect('login_portal')

    if request.method == "POST":
        course = get_object_or_404(Course, id=course_id)

        # 🛡️ TIME GATEKEEPER
        now = timezone.now()
        if now < course.start_time:
            messages.error(request, f"Examination begins at {course.start_time.strftime('%I:%M %p')}.")
            return redirect('login_portal')
        if now > course.end_time:
            messages.error(request, "Examination window has closed.")
            return redirect('login_portal')

        # 🛡️ Prevent duplicate submissions
        if StudentSubmission.objects.filter(index_number=index_number, course=course).exists():
            storage = get_messages(request)
            for _ in storage: pass
            messages.error(request, "You have already submitted this paper.")
            return redirect('login_portal')

        # 2. Randomization Engine
        raw_questions = list(Question.objects.filter(course=course))
        random.shuffle(raw_questions)

        questions_data = []
        question_ids = []  # For session tracking

        for q in raw_questions:
            question_ids.append(q.id)
            if q.q_type == 'MCQ':
                options = [{'val': 'A', 'text': q.option_a}, {'val': 'B', 'text': q.option_b},
                           {'val': 'C', 'text': q.option_c}, {'val': 'D', 'text': q.option_d}]
                random.shuffle(options)
                questions_data.append({'q': q, 'options': options})
            else:
                questions_data.append({'q': q, 'options': None})

        # 3. Store metadata in session for secure verification on submit
        request.session['exam_initialized'] = True
        request.session['exam_questions'] = question_ids

        context = {
            'course': course,
            'questions_data': questions_data,
            'student_name': student_name,
            'index_number': index_number,
            'duration_ms': course.duration_minutes * 60 * 1000
        }
        return render(request, 'quiz/exam.html', context)

    # If user
    return redirect('login_portal')


@csrf_exempt
def submit_quiz(request):
    if request.method == "POST":
        course_id = request.session.get('course_id')
        index_number = request.session.get('index_number')
        full_name = request.session.get('student_name')

        if not all([course_id, index_number]):
            return HttpResponse("Unauthorized", status=401)

        course_obj = get_object_or_404(Course, id=course_id)

        # 🛡️ DUPLICATE SUBMISSION CHECK
        if StudentSubmission.objects.filter(index_number=index_number, course=course_obj).exists():
            return HttpResponse("Form processing rejected: Duplicate submission.", status=403)

        # Process grading
        question_ids = request.session.get('exam_questions', [])
        all_questions = Question.objects.filter(id__in=question_ids) if question_ids else Question.objects.filter(
            course=course_obj)
        answers = {key: value for key, value in request.POST.items() if key.startswith('q')}

        correct_count = sum(1 for q in all_questions if
                            str(answers.get(f'q{q.id}', '')).strip().upper() == str(q.correct_answer).strip().upper())

        # Save submission
        StudentSubmission.objects.create(
            student_name=full_name,
            index_number=index_number,
            course=course_obj,
            submitted_answers=answers,
            score=float(correct_count)
        )

        request.session.flush()
        return render(request, 'quiz/submitted.html', {'score': int(correct_count)})

    return redirect('login_portal')

# ==========================================================
#             ADVANCED BULK QUESTIONS IMPORT ENGINE
# ==========================================================

def import_questions_page(request):
    """Renders the HTML upload interface page for lecturers."""
    if not request.user.is_staff:
        return redirect('login_portal')

    context = {
        'courses': Course.objects.all()
    }
    return render(request, 'admin/import_questions.html', context)


def import_questions_all_formats(request):
    """Processes document streams, routing parsing logic automatically based on extensions."""
    if request.method != "POST" or not request.FILES.get('uploaded_document'):
        return redirect('login_portal')

    course_id = request.POST.get('course_id')
    course_obj = get_object_or_404(Course, id=course_id)
    uploaded_file = request.FILES['uploaded_document']

    ext = os.path.splitext(uploaded_file.name)[1].lower()

    QTYPE_MAP = {
        'MCQ': 'MCQ', 'MULTIPLE CHOICE': 'MCQ', 'MULTIPLE_CHOICE': 'MCQ',
        'FITB': 'FITB', 'FILL IN THE BLANK': 'FITB', 'FILL_IN_THE_BLANK': 'FITB',
        'THEORY': 'THEORY', 'WRITTEN ESSAY': 'THEORY', 'ESSAY': 'THEORY'
    }

    try:
        if ext == '.xlsx':
            if not openpyxl:
                return HttpResponse("Server lacks Excel processor installation (openpyxl).", status=500)
            wb = openpyxl.load_workbook(uploaded_file, data_only=True)
            sheet = wb.active

            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row and len(row) >= 7 and row[0]:
                    raw_type = str(row[1]).strip().upper()
                    clean_type = QTYPE_MAP.get(raw_type, 'MCQ')

                    Question.objects.create(
                        course=course_obj,
                        text=str(row[0]).strip(),
                        q_type=clean_type,
                        option_a=str(row[2]).strip() if row[2] else "",
                        option_b=str(row[3]).strip() if row[3] else "",
                        option_c=str(row[4]).strip() if row[4] else "",
                        option_d=str(row[5]).strip() if row[5] else "",
                        correct_answer=str(row[6]).strip()
                    )

        elif ext == '.docx':
            doc = Document(uploaded_file)
            if doc.tables:
                table = doc.tables[0]
                for i, row in enumerate(table.rows):
                    if i == 0:
                        continue  # Skip headers

                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 7 and cells[0]:
                        raw_type = cells[1].strip().upper()
                        clean_type = QTYPE_MAP.get(raw_type, 'MCQ')

                        Question.objects.create(
                            course=course_obj,
                            text=cells[0], q_type=clean_type,
                            option_a=cells[2], option_b=cells[3], option_c=cells[4], option_d=cells[5],
                            correct_answer=cells[6]
                        )
            else:
                return HttpResponse("Word document parsing error: Could not find structured table grid layout.",
                                    status=400)

        elif ext == '.pdf':
            reader = PdfReader(uploaded_file)
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"

            lines = [line.strip() for line in full_text.split('\n') if line.strip()]
            for line in lines:
                parts = line.split(',')
                if len(parts) >= 7 and not parts[0].lower().startswith('text'):
                    raw_type = parts[1].strip().upper()
                    clean_type = QTYPE_MAP.get(raw_type, 'MCQ')

                    Question.objects.create(
                        course=course_obj,
                        text=parts[0].strip(), q_type=clean_type,
                        option_a=parts[2].strip(), option_b=parts[3].strip(),
                        option_c=parts[4].strip(), option_d=parts[5].strip(),
                        correct_answer=parts[6].strip()
                    )

        elif ext == '.csv':
            try:
                decoded_file = uploaded_file.read().decode('utf-8-sig')
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                decoded_file = uploaded_file.read().decode('latin-1')

            io_string = io.StringIO(decoded_file)
            reader = csv.reader(io_string)
            next(reader)  # Skip headers

            for row in reader:
                if len(row) >= 7 and row[0]:
                    raw_type = row[1].strip().upper()
                    clean_type = QTYPE_MAP.get(raw_type, 'MCQ')

                    Question.objects.create(
                        course=course_obj,
                        text=row[0].strip(), q_type=clean_type,
                        option_a=row[2].strip(), option_b=row[3].strip(),
                        option_c=row[4].strip(), option_d=row[5].strip(),
                        correct_answer=row[6].strip()
                    )
        else:
            return HttpResponse("Unsupported file format exception.", status=400)

        return redirect('/admin/quiz/question/')

    except Exception as e:
        return HttpResponse(f"System parsing exception caught during data synchronization: {str(e)}", status=500)


# ==========================================================
#             BLUEPRINT TEMPLATE DOWNLOAD ENDPOINTS
# ==========================================================

def download_excel_template(request):
    if not request.user.is_staff:
        return redirect('login_portal')
    if not openpyxl:
        return HttpResponse("Excel layout components missing on environment runtime.", status=500)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Question Import Template"

    headers = ['text', 'q_type', 'option_a', 'option_b', 'option_c', 'option_d', 'correct_answer']
    ws.append(headers)

    sample_row = [
        "Example Question: What language is Django written in?",
        "Multiple Choice", "Java", "Python", "C++", "PHP", "B"
    ]
    ws.append(sample_row)

    for cell in ws[1]:
        cell.font = openpyxl.styles.Font(bold=True)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="question_import_blueprint.xlsx"'
    wb.save(response)
    return response


def download_word_template(request):
    if not request.user.is_staff:
        return redirect('login_portal')

    doc = Document()
    doc.add_heading('Question Import Blueprint Grid', level=1)
    doc.add_paragraph('Fill your data inside the structured table layout below. Do not delete the top headers row.')

    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'

    headers = ['text', 'q_type', 'option_a', 'option_b', 'option_c', 'option_d', 'correct_answer']
    for idx, name in enumerate(headers):
        table.cell(0, idx).text = name
        table.cell(0, idx).paragraphs[0].runs[0].font.bold = True

    sample_values = ["Sample Target?", "Multiple Choice", "Opt A", "Opt B", "Opt C", "Opt D", "A"]
    for idx, value in enumerate(sample_values):
        table.cell(1, idx).text = value

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="question_import_blueprint.docx"'
    doc.save(response)
    return response




# ==========================================================
#         🔥 ALLOWED STUDENTS BULK ROSTER UPLOADER
# ==========================================================
def staff_member_required(view_func):
    return user_passes_test(lambda u: u.is_staff)(view_func)


@staff_member_required
def upload_allowed_students(request):
    """Processes incoming class rosters and safely registers students across multiple courses."""
    if request.method == "POST":
        course_id = request.POST.get("course_id")
        file_ref = request.FILES.get("student_file")

        if not course_id or not file_ref:
            messages.error(request, "Please select a valid course and attach a roster file.")
            return redirect("admin:quiz_allowedstudent_changelist")

        try:
            course = Course.objects.get(id=course_id)
            filename = file_ref.name.lower()
            students_to_create = []
            duplicate_count = 0
            success_count = 0

            # --- PROCESS EXCEL (.XLSX) ROSTERS ---
            if filename.endswith(".xlsx"):
                wb = openpyxl.load_workbook(file_ref, data_only=True)
                sheet = wb.active

                for row in sheet.iter_rows(min_row=2, values_only=True):
                    if not row or not row[0]:
                        continue

                    idx = str(row[0]).strip().upper()
                    name = str(row[1]).strip() if len(row) > 1 and row[1] else "Registered Student"

                    # 🛡️ WORKAROUND FILTER:
                    if AllowedStudent.objects.filter(index_number=idx, course=course).exists():
                        duplicate_count += 1
                        continue

                    students_to_create.append(AllowedStudent(course=course, index_number=idx, full_name=name))

            # --- PROCESS CSV (.CSV) ROSTERS ---
            elif filename.endswith(".csv"):
                try:
                    decoded_file = file_ref.read().decode('utf-8-sig').splitlines()
                except UnicodeDecodeError:
                    file_ref.seek(0)
                    decoded_file = file_ref.read().decode('latin-1').splitlines()

                reader = csv.reader(decoded_file)
                next(reader, None)

                for row in reader:
                    if not row or not row[0]:
                        continue

                    idx = str(row[0]).strip().upper()
                    name = str(row[1]).strip() if len(row) > 1 and row[1] else "Registered Student"

                    # 🛡️ WORKAROUND FILTER:
                    if AllowedStudent.objects.filter(index_number=idx, course=course).exists():
                        duplicate_count += 1
                        continue

                    students_to_create.append(AllowedStudent(course=course, index_number=idx, full_name=name))

            # Bulk insert the valid entries
            if students_to_create:
                AllowedStudent.objects.bulk_create(students_to_create)
                success_count = len(students_to_create)

            messages.success(request,
                             f"Roster processed: {success_count} students added, {duplicate_count} duplicates skipped.")

        except Exception as e:
            messages.error(request, f"Error processing file layout: {str(e)}")

    return redirect("admin:quiz_allowedstudent_changelist")
